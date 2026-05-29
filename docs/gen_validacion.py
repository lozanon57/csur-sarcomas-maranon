"""
Genera el documento formal de Política de Validación de Contenidos
para la CSUR Sarcomas App - HGUGM.

Uso: python3 docs/gen_validacion.py
Salida: docs/CSUR_Sarcomas_App_Validacion_Contenidos.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT = "docs/CSUR_Sarcomas_App_Validacion_Contenidos.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, 12), bold=True,
             color=(27, 66, 128) if level == 1 else (50, 50, 50))
    return p


def body(doc, text, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    set_font(run, italic=italic)
    return p


def bullet(doc, text, indent_cm=1.0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    return p


def add_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, size=10, bold=(i == 0))


# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Cover ────────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("POLÍTICA DE CONTENIDOS Y METODOLOGÍA DE VALIDACIÓN")
set_font(run, size=18, bold=True, color=(27, 66, 128))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("CSUR Sarcomas App")
set_font(run2, size=14, bold=True, color=(60, 60, 60))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(
    "Hospital General Universitario Gregorio Marañón\n"
    "Comité Multidisciplinar CSUR de Sarcomas y Tumores Musculoesqueléticos"
)
set_font(run3, size=12, italic=True)

doc.add_paragraph()

# Metadata table
meta = doc.add_table(rows=0, cols=2)
meta.style = "Table Grid"
meta.columns[0].width = Cm(5)
meta.columns[1].width = Cm(10)

data = [
    ("Documento",      "Política de Validación de Contenidos Clínicos"),
    ("Aplicación",     "CSUR Sarcomas App v1.1.0"),
    ("Centro",         "H.G.U. Gregorio Marañón — Madrid, España"),
    ("Responsable",    "Dr. Pablo Lozano Lominchar — Cirujano Oncológico"),
    ("Cargo",          "CSUR Sarcomas HGUGM · MSKCC Fellow 2024–25"),
    ("ORCID",          "0000-0002-5413-8449"),
    ("Fecha emisión",  "Mayo 2026"),
    ("Versión doc",    "1.0"),
    ("Clasificación",  "USO EXCLUSIVO CSUR — Herramienta educativa interna"),
]
for label, value in data:
    add_row(meta, label, value)

doc.add_page_break()

# ── 1. Objeto y alcance ──────────────────────────────────────────────────────

heading(doc, "1. Objeto y Alcance")
body(doc,
    "El presente documento describe la política de contenidos, las fuentes primarias "
    "utilizadas, el proceso de curación y revisión clínica, y las limitaciones de la "
    "aplicación CSUR Sarcomas App (en adelante, 'la aplicación'), desarrollada para "
    "uso exclusivo de los miembros del Comité Multidisciplinar CSUR de Sarcomas y "
    "Tumores Musculoesqueléticos del Hospital General Universitario Gregorio Marañón.")
body(doc,
    "La aplicación es una herramienta de apoyo educativo y de consulta clínica rápida. "
    "No genera diagnósticos ni recomendaciones terapéuticas de forma autónoma. "
    "Todas las decisiones clínicas sobre pacientes individuales deben adoptarse "
    "en el seno del comité multidisciplinar y de acuerdo con los protocolos "
    "institucionales vigentes.")
divider(doc)

# ── 2. Responsable ───────────────────────────────────────────────────────────

heading(doc, "2. Responsable de Contenidos y Validador Clínico")
body(doc,
    "El responsable de la elaboración, curación y validación de todos los contenidos "
    "clínicos es:")
doc.add_paragraph()

info_table = doc.add_table(rows=0, cols=2)
info_table.style = "Table Grid"
info_table.columns[0].width = Cm(5)
info_table.columns[1].width = Cm(10)

for label, value in [
    ("Nombre",        "Dr. Pablo Lozano Lominchar"),
    ("Especialidad",  "Cirugía General y del Aparato Digestivo"),
    ("Titulación",    "MD, PhD (UCM) — Sobresaliente Cum Laude"),
    ("Certificación", "European Board of Peritoneal Surface Malignancy (EBPSM)"),
    ("Fellowship",    "International Surgical Oncologist — Memorial Sloan Kettering "
                      "Cancer Center (MSKCC), Nueva York 2024–2025 (Peritoneal Service / "
                      "Sarcoma Service)"),
    ("Cargo",         "Cirujano Oncológico, Unidad de Cirugía de EMP, Sarcomas y Tumores "
                      "Pélvicos Avanzados — HGUGM"),
    ("Docencia",      "Profesor Asociado en Patología Quirúrgica, UCM"),
    ("Sociedad",      "Vocal Sección Sarcomas, Asociación Española de Cirujanos (AEC)"),
    ("ORCID",         "https://orcid.org/0000-0002-5413-8449"),
]:
    add_row(info_table, label, value)

doc.add_paragraph()
body(doc,
    "La idoneidad del responsable como validador clínico queda acreditada por su "
    "formación específica en sarcomas de partes blandas, óseos y abdominales en "
    "uno de los centros de referencia mundial (MSKCC), así como por su actividad "
    "asistencial e investigadora en el CSUR acreditado del HGUGM.",
    italic=True)
divider(doc)

# ── 3. Fuentes primarias ─────────────────────────────────────────────────────

heading(doc, "3. Fuentes Primarias de Contenido")
body(doc,
    "Todo el contenido clínico (descripciones histológicas, criterios diagnósticos, "
    "estadificación, algoritmos de tratamiento y seguimiento) ha sido extraído "
    "directamente de las siguientes guías clínicas y fuentes de evidencia:")

sources = [
    ("NCCN Clinical Practice Guidelines in Oncology",
     "v1.2025 — Soft Tissue Sarcoma; Bone Cancer; Gastrointestinal Stromal Tumors (GIST). "
     "National Comprehensive Cancer Network. https://www.nccn.org"),
    ("ESMO Clinical Practice Guidelines",
     "Soft Tissue and Visceral Sarcomas: ESMO–EURACAN–GENTURIS Clinical Practice "
     "Guidelines for diagnosis, treatment and follow-up. Annals of Oncology, 2025."),
    ("WHO Classification of Tumours — 5ª edición",
     "Soft Tissue and Bone Tumours. IARC Press, Lyon, 2020. ISBN 978-92-832-4502-5."),
    ("AJCC Cancer Staging Manual — 8ª edición",
     "American Joint Committee on Cancer. Amin MB et al. (eds.). Springer, 2017. "
     "Capítulos: Soft Tissue Sarcoma of the Trunk and Extremities; Retroperitoneal "
     "Sarcoma; Bone."),
    ("GEIS — Grupo Español de Investigación en Sarcomas",
     "Guías clínicas y documentos de consenso 2023–2025. www.geis.es"),
    ("Literatura científica indexada",
     "PubMed/MEDLINE. Revistas de referencia: Annals of Surgical Oncology, "
     "Lancet Oncology, Journal of Clinical Oncology, Annals of Oncology, "
     "Cancers (MDPI), British Journal of Surgery, European Journal of Cancer."),
    ("Ensayos clínicos históricos (Landmark Trials)",
     "24 ensayos seleccionados por su impacto demostrado en la práctica clínica "
     "estándar. Fuentes: PubMed, ClinicalTrials.gov. Nivel de evidencia asignado "
     "según escala Oxford CEBM (Ia–IV)."),
]
for i, (src, desc) in enumerate(sources, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(5)
    run_n = p.add_run(f"[{i}] {src}. ")
    set_font(run_n, size=11, bold=True)
    run_d = p.add_run(desc)
    set_font(run_d, size=11)

divider(doc)

# ── 4. Proceso de curación ───────────────────────────────────────────────────

heading(doc, "4. Proceso de Curación y Revisión Clínica")

heading(doc, "4.1 Elaboración del contenido", level=2)
bullet(doc, "El contenido de cada entidad nosológica (>30 sarcomas) fue elaborado "
       "directamente por el responsable a partir de las fuentes primarias listadas "
       "en el apartado 3, sin intermediarios ni resúmenes de terceros.")
bullet(doc, "Los datos de epidemiología, histología, inmunohistoquímica, marcadores "
       "moleculares, estadificación, tratamiento y seguimiento se contrastaron "
       "simultáneamente en NCCN 2025 y ESMO 2025.")
bullet(doc, "Las recomendaciones de tratamiento de primera línea reflejan el "
       "estándar de cuidado (Standard of Care) establecido en las guías más recientes "
       "en el momento de la elaboración (Mayo 2026).")

heading(doc, "4.2 Algoritmos de decisión clínica", level=2)
bullet(doc, "Los 6 algoritmos clínicos (diagnóstico, tratamiento y seguimiento) "
       "son árboles de decisión deterministas diseñados manualmente por el responsable "
       "basándose en los flujogramas de NCCN v1.2025 y ESMO 2025.")
bullet(doc, "Cada nodo del algoritmo cita explícitamente la fuente y versión de la "
       "guía en la que se basa.")
bullet(doc, "La aplicación NO utiliza inteligencia artificial generativa para producir "
       "ni modificar recomendaciones clínicas.")

heading(doc, "4.3 Ensayos históricos (Landmark Trials)", level=2)
bullet(doc, "La selección de 24 ensayos históricos se realizó conforme a criterios "
       "de impacto documentado en la práctica clínica estándar (cambio en guías, "
       "aprobaciones regulatorias, cambio en primera línea de tratamiento).")
bullet(doc, "A cada ensayo se asigna explícitamente un nivel de evidencia según la "
       "Oxford Centre for Evidence-Based Medicine (CEBM) 2011.")

heading(doc, "4.4 Revisión y actualización", level=2)
bullet(doc, "El contenido fue revisado íntegramente por el responsable en Mayo 2026 "
       "antes del lanzamiento de la versión v1.1.0.")
bullet(doc, "Se prevé revisión periódica coincidiendo con la publicación de nuevas "
       "versiones de las guías NCCN y ESMO (generalmente anual).")
bullet(doc, "Los cambios mayores se documentarán en el registro de versiones "
       "accesible desde la sección 'Acerca de' de la propia aplicación.")
divider(doc)

# ── 5. Limitaciones ──────────────────────────────────────────────────────────

heading(doc, "5. Limitaciones y Declaración de Responsabilidad")
body(doc,
    "La aplicación CSUR Sarcomas App es una herramienta de apoyo educativo. "
    "Las siguientes limitaciones deben tenerse en cuenta en su uso:")

limitations = [
    "El contenido tiene carácter exclusivamente educativo para los profesionales "
    "sanitarios miembros del CSUR. No constituye consejo médico individualizado.",
    "Las recomendaciones clínicas contenidas en la aplicación reflejan el estándar "
    "de cuidado general definido en las guías. La decisión final sobre cada paciente "
    "debe adoptarse siempre en el contexto del equipo multidisciplinar y de acuerdo "
    "con los protocolos institucionales del HGUGM.",
    "Las guías clínicas se actualizan con periodicidad variable. Puede existir un "
    "desfase entre la publicación de nuevas recomendaciones y su incorporación a la "
    "aplicación.",
    "Los ensayos clínicos activos se presentan con fines informativos. La elegibilidad "
    "real de cada paciente debe verificarse directamente en la ficha oficial del ensayo "
    "(ClinicalTrials.gov) y con el equipo investigador.",
    "La aplicación no sustituye la consulta directa a las guías originales ni al juicio "
    "clínico experto.",
]
for lim in limitations:
    bullet(doc, lim)

doc.add_paragraph()
body(doc,
    "El responsable de los contenidos no asume responsabilidad por decisiones clínicas "
    "adoptadas exclusivamente con base en la información contenida en la aplicación, "
    "sin contraste con las guías originales y sin deliberación en el comité "
    "multidisciplinar correspondiente.",
    italic=True)
divider(doc)

# ── 6. Control de versiones ──────────────────────────────────────────────────

heading(doc, "6. Control de Versiones")

ver_table = doc.add_table(rows=1, cols=4)
ver_table.style = "Table Grid"
headers = ["Versión", "Fecha", "Responsable", "Descripción del cambio"]
for i, h in enumerate(headers):
    cell = ver_table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, size=10, bold=True)

versions = [
    ("v1.0.0", "Marzo 2026",  "Dr. P. Lozano", "Versión inicial. 25 entidades, 5 algoritmos, buscador."),
    ("v1.1.0", "Mayo 2026",   "Dr. P. Lozano", "30+ entidades, 6 algoritmos, 24 ensayos históricos, "
                                                "perla docente diaria, sección Metodología y Validación."),
]
for ver, fecha, resp, desc in versions:
    row = ver_table.add_row()
    for i, val in enumerate([ver, fecha, resp, desc]):
        row.cells[i].text = val
        for run in row.cells[i].paragraphs[0].runs:
            set_font(run, size=10)

divider(doc)

# ── 7. Firma ─────────────────────────────────────────────────────────────────

heading(doc, "7. Firma y Fecha de Validación")
body(doc, "El responsable certifica que el contenido clínico de la aplicación "
     "CSUR Sarcomas App v1.1.0 ha sido elaborado conforme a las fuentes primarias "
     "descritas en este documento y revisado clínicamente en Mayo de 2026.")

doc.add_paragraph()
doc.add_paragraph()

sig = doc.add_paragraph()
sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = sig.add_run(
    "Dr. Pablo Lozano Lominchar\n"
    "MD, PhD · EBPSM · MSKCC Fellow 2024–25\n"
    "Cirujano Oncológico — CSUR Sarcomas\n"
    "Hospital General Universitario Gregorio Marañón · Madrid\n"
    f"Fecha: Mayo 2026\n"
    "ORCID: 0000-0002-5413-8449"
)
set_font(run, size=11)

doc.add_paragraph()
doc.add_paragraph()

line = doc.add_paragraph("_" * 50)
line.paragraph_format.space_after = Pt(2)
caption = doc.add_paragraph("Firma del responsable de contenidos")
caption.paragraph_format.space_before = Pt(0)
for run in caption.runs:
    set_font(run, size=10, italic=True)

# ── Footer note ──────────────────────────────────────────────────────────────

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.paragraph_format.space_before = Pt(20)
run_f = footer_p.add_run(
    "© 2026 Dr. Pablo Lozano Lominchar · CSUR Sarcomas HGUGM · Todos los derechos reservados. "
    "Documento generado automáticamente — CSUR Sarcomas App v1.1.0."
)
set_font(run_f, size=9, italic=True, color=(120, 120, 120))

# ── Save ─────────────────────────────────────────────────────────────────────

doc.save(OUT)
print(f"✓ Documento generado: {OUT}")
